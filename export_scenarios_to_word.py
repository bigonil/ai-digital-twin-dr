#!/usr/bin/env python3
"""
Export MCP Server usage scenarios to professional Word document.
Requires: python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_scenario(doc, title, intro, steps):
    """Add a scenario section."""
    doc.add_heading(title, level=2)
    if intro:
        p = doc.add_paragraph(intro)
        p.style = 'List Bullet'
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(step, style='List Number')

def shade_cell(cell, color):
    """Shade a table cell with color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_document():
    """Create the Word document."""
    doc = Document()

    # Title
    title = doc.add_heading('AI Digital Twin - MCP Server', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Scenari Realistici di Utilizzo e Integrazioni')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.italic = True
    subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacer

    # Scenario 1
    add_heading(doc, '📊 Scenario 1: Incident Response in Real-Time', level=1)
    intro = '🚨 Allerta: db-001-primary è down!'
    steps = [
        'Chiedi al MCP: simulate_disaster(db-001-primary) → Vedi subito che 16 nodi saranno colpiti',
        'Chiedi: get_recovery_plan(db-001-primary) → Hai il runbook step-by-step già disponibile',
        'Chiedi: get_simulation_timeline(simulation_id) → Vedi l\'ordine esatto di fallimento per coordinare il team',
    ]
    add_scenario(doc, '', intro, steps)

    doc.add_paragraph()  # Spacer

    # Scenario 2
    add_heading(doc, '📊 Scenario 2: Disaster Recovery Planning', level=1)
    intro = 'Plan: Migliorare RTO del 50% entro Q2'
    steps = [
        'Simula tutti i nodi critici: simulate_disaster(db-001-primary), simulate_disaster(api-001), simulate_disaster(lb-001)',
        'Analizza: analyze_cascading_failure() per ogni simulation → Identifica bottleneck (es: 75min per AWS Backup)',
        'Proponi: "Aggiungiamo multi-region replica?" → Re-simula con nuova topologia',
    ]
    add_scenario(doc, '', intro, steps)

    doc.add_paragraph()  # Spacer

    # Scenario 3
    add_heading(doc, '📊 Scenario 3: Infrastructure Compliance & Audit', level=1)
    intro = 'Auditoria: "Verifichiamo che live infrastructure = IaC"'
    steps = [
        'Chiedi: check_drift() → Scopri risorse lanciate fuori da Terraform',
        'Risolvi drift: aggiorna Terraform files',
        'Re-esegui: check_drift() → zero drift ✅',
    ]
    add_scenario(doc, '', intro, steps)

    doc.add_paragraph()  # Spacer

    # Scenario 4
    add_heading(doc, '📊 Scenario 4: What-If Analysis per Decisioni Architetturali', level=1)
    intro = 'Domanda: "Vale la pena aggiungere un replica in secondary AZ?"'
    steps = [
        'simulate_disaster(db-001-primary, depth=5) → Baseline: worst-case RTO 120 min',
        'Modifica topologia: aggiungi db-003-secondary in us-west-2',
        'simulate_disaster(db-001-primary) → RTO 30 min ✅ → Decisione: ROI chiaro → implementa la modifica',
    ]
    add_scenario(doc, '', intro, steps)

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # Integrazioni
    add_heading(doc, '💡 Integrazioni Possibili', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Integrazione'
    hdr_cells[1].text = 'Usa MCP Per'

    # Shade header
    for cell in hdr_cells:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    integrations = [
        ('ChatOps (Slack, Teams)', 'Simula disaster on-demand, notifiche risultati'),
        ('IaC Pipeline (Terraform/Bicep)', 'check_drift() pre-merge, validation'),
        ('Incident Management (PagerDuty, Opsgenie)', 'Auto-trigger get_recovery_plan() su alert'),
        ('Cost Analysis', 'Analizza RTO trade-offs vs infrastructure cost'),
        ('Chaos Engineering (Gremlin, Chaos Toolkit)', 'Valida predictions vs real chaos tests'),
        ('Knowledge Base (Confluence, Notion)', 'Auto-export recovery plans a runbook'),
    ]

    for integration, usage in integrations:
        row_cells = table.add_row().cells
        row_cells[0].text = integration
        row_cells[1].text = usage

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # Cosa Manca
    add_heading(doc, '🔧 Cosa Manca Oggi', level=1)
    intro_text = doc.add_paragraph('Per una soluzione completamente end-to-end, mancano:')

    missing = [
        'Live Monitoring Integration — RTO/RPO ora sono statici, dovrebbero venire da VictoriaMetrics',
        'Multi-cloud Inference — Topology è solo AWS, dovrebbe supportare Azure + GCP',
        'Real-Time Graph Sync — Neo4j si aggiorna solo manualmente con Terraform parser, dovrebbe auto-aggiornarsi da AWS APIs',
        'Chaos Validation — Simulazioni vs real test results — nessun confronto automatico',
    ]

    for item in missing:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()  # Spacer

    # Next Steps
    add_heading(doc, '🚀 Prossimi Passi', level=1)
    intro_text = doc.add_paragraph('Quale priorità implementare subito?')

    options = [
        'Live Monitoring — Connetti VictoriaMetrics per RTO/RPO real-time',
        'Multi-cloud — Aggiungi Azure SQL + GCP Cloud SQL alla topologia',
        'Auto-Sync — API credentials per auto-aggiornare grafo da AWS/Azure/GCP',
    ]

    for option in options:
        doc.add_paragraph(option, style='List Number')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('—')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = doc.add_paragraph('AI Digital Twin - MCP Server | Professional Disaster Recovery Planning')
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(10)
    footer_text.runs[0].font.italic = True
    footer_text.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = r'C:\Users\luca.bigoni\ai-digital-twin-dr\MCP_Server_Scenarios.docx'
    doc.save(output_path)
    print(f'[OK] Documento creato: {output_path}')
